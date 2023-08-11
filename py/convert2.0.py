import argparse
# 传参：filePath
from docx import Document
# , default = r'C:\Users\Dell\Desktop\TextConvert\加涅信息加工系统是用于处理什么类型的数据的？.txt'
parser = argparse.ArgumentParser(description='manual to this script')
parser.add_argument('--fileName', type=str, default=r'D:\lams office\test\加涅信息加工系统是用于处理什么类型的数据的？.txt')
parser.add_argument('--targetName', type=str, default=r'D:\lams office\test\加涅信息加工系统是用于处理什么类型的数据的？.docx')

args = parser.parse_args()

def convertext(fileName,targetName):  #path指的是需要处理文档的路径，name指的是新生成文档的名称，houzhui指的是新生成文档的形式
    doc = Document()
    # doc.save('converted.docx')
    file = open(fileName, 'r', encoding='utf-8')
    text = file.read()
    # text='问题1: 当前哪些省市的干部群众仍在与暴雨洪灾作顽强斗争？\n选项：\nA) 京津冀黑吉\nB) 沪浙皖赣\nC) 鲁豫陕川\nD) 粤桂琼闽\n答案: A) 京津冀黑吉\n\n问题2: 保定市县乡干部抵达\"孤岛\"后，与汤家庄村群众共同采取了什么措施来救援被困村民？\n选项：\nA) 建立临时救援基地\nB) 借助绳索将被困村民背出\nC) 泡制救生圈漂流出村落\nD) 乘坐橡皮艇破浪前行\n答案: B) 借助绳索将被困村民背出\n\n问题3: 受困期间，K1178次列车上采取了哪些措施来帮助旅客？\n选项：\nA) 成立临时救援基地\nB) 成立临时党支部，征集志愿者\nC) 联系外部救援物资\nD) 安排大巴车接送旅客离开\n答案: B) 成立临时党支部，征集志愿者\n\n问题4: 在K1178次列车上，旅客自发将哪些药品交给列车集中管理，供大家使用？\n选项：\nA) 降压药、消炎药\nB) 止痛药、感冒药\nC) 止血药、救心丸\nD) 退烧药、胃药\n答案: A) 降压药、消炎药'
    # print(text)

    list1 = text.split("\n")
    # print(list1)

    ans = ""
    # answer_re=list1[5][:2]
    for index, item in enumerate(list1):
        # print(item[:2])
        # print(item[:2] == '答案')
        if item:
            if item[:7] == "Answer:":
                ans += item
            elif item[:2] not in ["A)", "B)", "C)", "D)","A.","B.","C.","D.","a)","b)","c)","d)","a.","b.","c.","d."] and item[:4] not in ["正确答案","答案：A", "答案：B", "答案：C", "答案：D","答案：a","答案：b","答案：c","答案：d"]:
                ans += "Question: "
                ans += item
            elif item[:2] == "答案" or item[:4] == "正确答案":
                alb = item[-1]
                if item[-1] in ["A","B","C","D"]:
                    alb = chr(ord(alb) + 32)
                ans += "Answer:"
                ans += alb
            elif item[:2] in ["A)", "B)", "C)", "D)","A.","B.","C.","D.","a)","b)","c)","d)","a.","b.","c.","d."] :
                if item[0] in ["A","B","C","D"]:
                    ans += chr(ord(item[0]) + 32)
                else:
                    ans+=item[0]
                ans += item[1:]

            doc.add_paragraph(ans)
            ans=""
    #    ans += "\n"
		
   # para = doc.add_paragraph(ans)
    doc.save(targetName)
    # doc.save('abc5.docx')


convertext(args.fileName,args.targetName)

# 运行合格：python convert.py --filePath 需要被转换的txt文档的全路径
